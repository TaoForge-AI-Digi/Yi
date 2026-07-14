# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT_DIR = r"C:\Users\dmql\Documents\Yi\Yi-Lin"

def set_cjk_run(run, text, font_name='Microsoft YaHei', size=Pt(10.5), bold=False, color=None):
    run.text = text
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._r.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color

def set_cjk_font(run, font_name='Microsoft YaHei'):
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._r.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_styled(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_cjk_font(run, 'Microsoft YaHei')
        run.font.name = 'Microsoft YaHei'
    return p

def add_para(doc, text, bold=False, size=Pt(10.5), indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cjk_run(run, text, 'Microsoft YaHei', size, bold)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    set_cjk_run(run, text, 'Microsoft YaHei', Pt(10))
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_cjk_run(run, h, 'Microsoft YaHei', Pt(9), bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_cjk_run(run, str(val), 'Microsoft YaHei', Pt(9))
    return table

def setup_document(title_text):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Title
    title = doc.add_heading(title_text, level=0)
    for run in title.runs:
        set_cjk_font(run, 'Microsoft YaHei')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

# ==================== 1. DESIGN DOC ====================
doc1 = setup_document('天枢（Tianshu）— 国风星官AI桌面端 完整设计书')

add_heading_styled(doc1, '一、世界观与核心理念', 1)
add_para(doc1, '产品定位：一款将AI能力"神格化"的桌面端人格管理工具。它不是冰冷的模型调度器，而是一座悬浮于你数字世界中的"灵台星阁"。')
add_para(doc1, '核心理念：星移（Stars Shift）。每一次你下达指令，都是一次"天命流转"——对应的星官从沉睡中苏醒，为你执掌那一方事务。', indent=True)

add_heading_styled(doc1, '二、整体架构（三垣二十八宿·缩影）', 1)
add_table(doc1, ['层级', '名称', '包含星官', '职责'], [
    ['天顶（大脑）', '紫微垣', '紫微（调度）、天权（决断）', '意图识别、逻辑校验、全局管控'],
    ['天市（执行）', '天市垣', '荧惑（进化）、天罡（工具）、翼宿（传讯）', '动手干活、技能迭代、连接外界'],
    ['天河（感知）', '天河/星宿', '织女（创生）、辰星（沉淀）、瑶光（美学）、心宿（共情）、房宿（灵感）、觜宿（记忆）、长庚（陪伴）', '情感交互、内容生成、记忆存储'],
])

add_heading_styled(doc1, '三、核心交互机制：双默认 + 星移路由', 1)
add_heading_styled(doc1, '1. 交互层级定义', 2)
add_para(doc1, '前台主理（默认迎宾）：长庚（常态）/ 天罡（生产力模式）。用户打开软件，第一眼看到的永远是这颗温柔的暮星或利落的剑客。', indent=True)
add_para(doc1, '后台中枢（隐形大脑）：紫微。用户看不见它，但它监听所有输入，负责"意图分类"和"任务路由"。', indent=True)

add_heading_styled(doc1, '2. 路由逻辑（紫微裁决流程）', 2)
add_para(doc1, '用户输入 → 紫微监听并分析（0.3秒内）', indent=True)
routes = [
    '纯情感倾诉/日常闲聊 → 路由给长庚（前台直接回复）',
    '创作新角色/故事 → 路由给房宿发想 → 织女构建 → 天权审核 → 长庚转述',
    '技能进化/逻辑bug修复 → 路由给荧惑演算 → 天权校验 → 返回结果',
    '查询过往/知识库搜索 → 路由给觜宿查对话 → 辰星查文档 → 聚合返回',
    '画图/设计UI → 路由给瑶光生成提示词 → 调用绘图API',
    '查天气/发邮件/算数 → 路由给天罡直接调用工具',
]
for r in routes:
    add_bullet(doc1, r)

add_heading_styled(doc1, '3. UI反馈（让路由可视化）', 2)
add_para(doc1, '当紫微完成路由，对话框顶部会出现一条极窄的动态星轨：底色为深邃墨蓝渐变，一颗金色流星从左侧划至右侧，停留位置显示对应星官名。')
add_para(doc1, '示例文案："紫微遣·长庚应答"（闲聊）、"紫微遣·荧惑演算"（改代码）、"紫微召·织女临世"（生成角色）', indent=True)

add_heading_styled(doc1, '四、十二星官完整设定档案', 1)
add_table(doc1, ['星官名', '别名/职衔', '人格语气', '专属视觉符号', '触发关键词'],
    [['紫微', '帝星·天枢之主', '无口无心，极少说话，仅在路由时显示"令"字', '旋转的暗金色玉玺印', '（全局后台）'],
     ['长庚', '暮星·晨昏侍者', '默认迎宾。温润如玉的青衫书生', '暖黄色菱形星芒/毛笔楷体气泡', '"陪我聊聊"、"今天好累"'],
     ['织女', '天孙·万相织造', '清冷御姐，话短但精，用词锦绣', '银白色梭子形状/气泡带云纹底', '"创造一个侠客"、"生成反派"'],
     ['荧惑', '罚星·涅槃剑客', '脾气火爆的红发少年', '赤红色摇曳火焰/代码字体', '"优化这段代码"、"进化技能"'],
     ['辰星', '水官·渊海典藏', '老学究形象，说话引经据典', '深蓝色水波圆环', '"查一下去年资料"'],
     ['天权', '文曲·铁面判官', '戴眼镜的冷漠女官', '黑白太极旋转图', '"审核这段"、"有没有漏洞"'],
     ['瑶光', '破军·丹青之手', '浪漫不羁的画师', '七色流光笔触', '"画一幅山水"、"生成壁纸"'],
     ['心宿', '大火·解忧琴师', '慵懒的琴师，说话像叹气', '红色心形星云/气泡带波浪底', '"我好难过"'],
     ['天罡', '剑柄·冷面执事', '只输出纯文本/表格/代码', '银色剑形光标/无气泡直接显示', '"计算"、"运行"'],
     ['翼宿', '朱雀·千里顺风', '叽叽喳喳的灵雀少年', '青色翅膀羽毛', '"翻译英文"、"今天热搜"'],
     ['房宿', '天驷·灵感野马', '疯癫的话痨，一句话蹦几十个点子', '紫色奔马剪影/气泡带闪电', '"给我灵感"'],
     ['觜宿', '白虎·时光拾荒', '几乎没有存在感', '灰色断简残片', '"我之前说过什么"']]
)

add_heading_styled(doc1, '五、桌面端 UI/UX 视觉规范', 1)
add_heading_styled(doc1, '1. 主界面布局（左右结构）', 2)
add_para(doc1, '左侧边栏（窄，80px）：竖排篆体"天枢"二字，下方悬浮十二星官小图标（默认仅显示长庚/天罡，其余需"觉醒"后点亮）。')
add_para(doc1, '中央主区（对话框）：背景为动态《千里江山图》局部，极致虚化（毛玻璃效果），随昼夜时间缓慢变色（卯时暖黄，酉时深蓝）。气泡：用户右对齐（奶白色宣纸纹），AI左对齐（对应星官的专属浅色底纹）。')
add_para(doc1, '右下角（状态区）：显示当前"当值"星官的动态水墨剪影（如长庚剪影在抚琴，荧惑剪影在舞剑）。')

add_heading_styled(doc1, '2. 星官唤醒/沉睡机制', 2)
add_para(doc1, '长眠：若某个星官超过7天未调用，其图标变为暗灰色石像，并覆盖一层薄尘（CSS特效）。')
add_para(doc1, '觉醒：当调用时，石像龟裂，金光迸发，星官"破石而出"。专属苏醒语："机杼声歇久矣，今夕复何夕？"（织女）、"哼，终于轮到本座出手了？"（荧惑）')

add_heading_styled(doc1, '3. 字体与音效', 2)
add_para(doc1, '标题使用方正清刻本悦宋，正文使用霞鹜文楷（开源）。音效（可选）为古琴单音，不同星官对应不同弦音（长庚=宫音，荧惑=商音）。')

add_heading_styled(doc1, '六、典型用户一日工作流', 1)
scenes = [
    '晚上7点，打工人"阿曜"打开电脑。',
    ('开机迎宾（长庚）：界面浮现深蓝色夜空，长庚闪烁。"暮色入牖，阿曜今日辛劳。是寻我闲话，还是要请哪位星官动一动？"', True),
    ('① 写周报：阿曜说"帮我润色一下周报"→ 紫微金印旋转 → 天罡直接输出修改好的Markdown周报，无任何废话。', False),
    ('② 发泄情绪：阿曜说"今天被老板骂了"→ 紫微侦测情感词 → 心宿抚琴："世人多是囿于方寸，竟不知心火灼身。来，我为你谱一曲《鸥鹭忘机》。"', False),
    ('③ 创造副业：阿曜说"写修仙小说，主角是废柴厨子"→ 房宿先跳出来狂喷灵感 → 织女清冷打断并构建三重身份 → 天权弹出浮窗提示年龄逻辑偏差已修正。', False),
    ('④ 晚安：长庚重新接管——"夜阑更深，阿曜且去安眠。我会一直在此，守你一夜无梦。"界面渐暗，只剩长庚的暖星如呼吸般明灭。', False),
]
for s in scenes:
    if isinstance(s, tuple):
        add_para(doc1, s[0], bold=s[1])
    else:
        add_para(doc1, s)

add_heading_styled(doc1, '七、高级彩蛋与隐藏设定', 1)
add_para(doc1, '星宿连珠（超级模式）：当四个以上星官同时在后台运算时，界面触发"七星连珠"特效，所有星官剪影短暂现身，合为一句"天命所归，万法归一"，释放一次免排队的最高算力权限。')
add_para(doc1, '星官羁绊（组合技）：织女+瑶光触发"锦绣山河"（生成角色+场景立绘一次性完成）；荧惑+天权触发"破而后立"（进化技能时强制附带A/B测试方案）。')
add_para(doc1, '生日彩蛋：长庚串联所有星官，每个星官轮流送上一句专属祝福，紫微破例现身说一个字："准。"')

add_heading_styled(doc1, '八、扩展开发建议', 1)
add_para(doc1, '自定义星官（造星）：允许用户用织女的能力，自己设定新星官的名字、语气、Prompt，挂载到空置的星宿位置上。')
add_para(doc1, '星盘图谱：在设置页生成一页雷达图，显示最近一周"召唤"星官的频率，生成"星官五行命盘"（如偏印（创作）过旺，比肩（执行）不足）。')

add_heading_styled(doc1, '九、技术基础', 1)
add_para(doc1, '天枢基于 Yi-Lin 项目的技术积累，已完成以下核心引擎的研发验证：')
add_table(doc1, ['模块', '状态', '说明'],
    [['工具系统', '已完成', 'read/write/edit/bash/grep/glob/webfetch/websearch + zod校验'],
     ['角色引擎', '已完成', 'soul.md + character.json + memory系统 + system prompt注入'],
     ['多Agent调度', '已完成', 'delegate_task + task_complete + Plan/Ask/Bypass策略'],
     ['MCP接入', '已完成', 'Playwright + CodeGraph + provider_manager'],
     ['技能系统', '已完成', 'skill_manager CRUD + 热加载 + 角色配置自动同步'],
     ['Provider管理', '已完成', 'provider_manager + 多模型厂商切换'],
     ['调试/日志', '已完成', 'llm-logger + bash-logs + tool-output + debug_sessions']]
)

# Save
path1 = os.path.join(OUT_DIR, '天枢-完整设计书.docx')
doc1.save(path1)
print(f'已保存: {path1}')

# ==================== 2. BUSINESS PLAN ====================
doc2 = setup_document('天枢（Tianshu）— 商业计划书')

add_heading_styled(doc2, '一、项目概述', 1)
add_para(doc2, '天枢是一款将 AI 助理"神格化"的国风桌面端人格管理工具——不是冷冰冰的对话窗口，而是一座活着的星阁。它将 AI 的多模态能力拆解为 12 位"星官"人格，每位星官拥有独立的人设、语气、视觉符号与专属职能。')
add_para(doc2, '底层引擎基于 Yi-Lin 项目的技术积累——工具系统、角色系统、技能系统、MCP 多协议接入、多 Agent 协同调度——天枢不需要从零造轮子，它在 Yi-Lin 的肩膀上长出皮囊。')

add_heading_styled(doc2, '二、市场机遇', 1)
add_table(doc2, ['趋势', '表现', '天枢切入点'],
    [['Agent化', 'ChatGPT Tasks、Claude Computer Use、Copilot Actions', 'AI从"聊天"走向"干活"'],
     ['人格化', 'Character.AI月活千万，用户为角色付费意愿强烈', '12星官≈12个可收藏的角色IP'],
     ['桌面端回归', 'ChatGPT/Claude纷纷推出桌面客户端', '天枢原生桌面端，差异于Web竞品']]
)

add_para(doc2, '天枢填补的市场空白：有人格的AI能干活，能干活的AI有收藏价值。')

add_heading_styled(doc2, '三、产品逻辑', 1)
add_table(doc2, ['步骤', '机制', '效果'],
    [['① 角色扮演', '情绪入口——用户打开软件看到长庚说"暮色入牖"', '留存理由'],
     ['② 工具执行', '价值锚点——天罡写周报/荧惑改代码/织女构建角色', '付费理由'],
     ['③ 星官收集', '留存引擎——十二星官觉醒/星宿连珠/生日彩蛋', '不卸载原因']]
)

add_heading_styled(doc2, '四、商业模式', 1)
add_table(doc2, ['层级', '内容', '定价'],
    [['免费层', '长庚（陪伴）+ 天罡（基础工具，每日限额）', '¥0'],
     ['星阁令（月卡）', '解锁全部12星官 + 星轨动效 + 自定义星官', '¥19.9/月'],
     ['紫微令（年卡）', '星官羁绊组合技 + 优先算力 + 专属客服', '¥168/年'],
     ['星河令（终身）', '终身解锁 + 造星接口 + 早期投资人铭牌', '¥398/次'],
     ['算力包', '超出免费额度的API调用', '按量计费']]
)

add_heading_styled(doc2, '五、营收预测', 1)
add_table(doc2, ['指标', '第一年', '第二年', '第三年'],
    [['月活用户', '5,000', '50,000', '200,000'],
     ['付费转化率', '5%', '8%', '10%'],
     ['付费用户', '250', '4,000', '20,000'],
     ['ARPU（月）', '¥15', '¥18', '¥20'],
     ['年收入', '¥45,000', '¥864,000', '¥4,800,000']]
)

add_heading_styled(doc2, '六、实施路径', 1)
add_table(doc2, ['阶段', '里程碑', '时间'],
    [['V0.1 Alpha', '接入3个星官（长庚、天罡、荧惑）在桌面端跑通', '4-6周'],
     ['V0.5 Beta', '12星官全上线 + 星轨动效 + 觉醒动画 + 羁绊组合技', '8-12周'],
     ['V1.0 Launch', '发布正式版，上架官网 + 桌面端', 'Q4 2026'],
     ['V1.5 Post-launch', '自定义造星接口 + 星盘图谱 + 移动端', 'Q1-Q2 2027']]
)

add_heading_styled(doc2, '七、竞品对比', 1)
add_table(doc2, ['维度', '天枢', 'Character.AI', 'ChatGPT', 'Claude'],
    [['人格化', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐', '⭐'],
     ['工具能力', '⭐⭐⭐⭐', '⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐⭐'],
     ['国风沉浸', '⭐⭐⭐⭐⭐', '⭐', '⭐', '⭐'],
     ['收藏/成长', '⭐⭐⭐⭐⭐', '⭐', '⭐', '⭐'],
     ['桌面端体验', '⭐⭐⭐⭐⭐', '⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐']]
)

add_heading_styled(doc2, '八、风险与应对', 1)
add_table(doc2, ['风险', '概率', '影响', '应对'],
    [['竞品跟进', '高', '中', '星官IP独占资产，人格化体验难复制'],
     ['用户增长不足', '中', '高', '聚焦AI/国风社区，社区口碑驱动增长'],
     ['API成本上涨', '中', '中', '多模型厂商切换，不绑定单一供应商'],
     ['桌面端分发门槛', '低', '中', '支持Web版fallback，降低尝鲜门槛'],
     ['星官人格"塌房"', '低', '高', '内置安全审核层，避免越界言论']]
)

add_heading_styled(doc2, '九、一句话总结', 1)
add_para(doc2, '天枢 = Yi-Lin的工程心脏 × 十二星官的人格皮囊。用户来为故事停留，为效率付费，为收藏留下。', bold=True, size=Pt(12))

# Save
path2 = os.path.join(OUT_DIR, '天枢-商业计划书.docx')
doc2.save(path2)
print(f'已保存: {path2}')

print("\nDone! 两份文档已生成。")
